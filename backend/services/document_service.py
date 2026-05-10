"""
Document processing service — loads, chunks, and prepares documents.
Supports PDF, DOCX, TXT, and plain text.
Lazy imports for heavy libraries (pypdf, python-docx).
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from models.schemas import DocumentMetadata

logger = logging.getLogger(__name__)


class ChunkedDocument:
    """A document split into chunks with metadata."""

    def __init__(
        self,
        chunks: List[str],
        metadata: DocumentMetadata,
        chunk_metadata: Optional[List[DocumentMetadata]] = None,
    ):
        self.chunks = chunks
        self.metadata = metadata
        self.chunk_metadata = chunk_metadata or [metadata] * len(chunks)

    def __len__(self) -> int:
        return len(self.chunks)

    def __getitem__(self, idx: int):
        return self.chunks[idx], self.chunk_metadata[idx]


class TextSplitter:
    """
    Recursive character text splitter with configurable chunk size and overlap.
    Mimics LangChain's RecursiveCharacterTextSplitter without dependency.
    """

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be less than chunk_size")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._separators = ["\n\n", "\n", ". ", "! ", "? ", ", ", " ", ""]

    def split_text(self, text: str) -> List[str]:
        """Split text into chunks recursively."""
        return self._split(text, self.chunk_size, self.chunk_overlap, self._separators)

    def _split(
        self,
        text: str,
        chunk_size: int,
        chunk_overlap: int,
        separators: List[str],
    ) -> List[str]:
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        chunks: List[str] = []
        current = text

        while current:
            if len(current) <= chunk_size:
                chunks.append(current)
                break

            # Find split point
            split_idx = self._find_split(current, chunk_size, separators)
            chunk = current[:split_idx].strip()
            if chunk:
                chunks.append(chunk)

            # Advance with overlap (ensure forward progress)
            if split_idx > chunk_overlap:
                overlap_start = split_idx - chunk_overlap
            else:
                overlap_start = split_idx  # no overlap if chunk is too small
            current = current[overlap_start:]

        return chunks

    def _find_split(self, text: str, max_len: int, separators: List[str]) -> int:
        """Find the best position to split within max_len."""
        if len(text) <= max_len:
            return len(text)

        # Try each separator from longest to shortest
        for sep in separators:
            if not sep:
                return max_len  # fallback: split at max_len
            pos = text.rfind(sep, 0, max_len)
            if pos > 0:
                return pos + len(sep)

        return max_len

    def split_documents(
        self,
        texts: List[str],
        metadatas: Optional[List[Dict[str, Any]]] = None,
    ) -> List[tuple[str, Dict[str, Any]]]:
        """Split multiple texts with optional metadata."""
        result: List[tuple[str, Dict[str, Any]]] = []
        for i, text in enumerate(texts):
            meta = metadatas[i] if metadatas and i < len(metadatas) else {}
            chunks = self.split_text(text)
            for chunk in chunks:
                result.append((chunk, meta))
        return result


class DocumentProcessor:
    """
    Loads documents from various formats and splits them into chunks.
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        self.splitter = TextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def load_and_chunk(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> ChunkedDocument:
        """
        Load a file and split into chunks.

        Args:
            file_path: Path to document (PDF, DOCX, TXT, MD).
            metadata: Additional metadata to attach.

        Returns:
            ChunkedDocument with chunks and metadata.

        Raises:
            ValueError: If file type is unsupported.
            FileNotFoundError: If file doesn't exist.
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Merge: explicit metadata overrides defaults, but avoid duplicate kwargs
        merged = {"source": path.name, "file_type": ext.lstrip(".")}
        if metadata:
            merged.update(metadata)

        base_meta = DocumentMetadata(**merged)

        if ext == ".pdf":
            return self._load_pdf(path, base_meta)
        elif ext in (".docx", ".doc"):
            return self._load_docx(path, base_meta)
        else:  # .txt, .md
            return self._load_text(path, base_meta)

    def chunk_text(
        self,
        text: str,
        metadata: Optional[DocumentMetadata] = None,
    ) -> ChunkedDocument:
        """
        Split raw text into chunks directly.

        Args:
            text: Raw text content.
            metadata: Optional metadata.

        Returns:
            ChunkedDocument.
        """
        meta = metadata or DocumentMetadata(source="manual")
        chunks = self.splitter.split_text(text)
        logger.debug("Split text into %d chunks", len(chunks))
        return ChunkedDocument(chunks=chunks, metadata=meta)

    def chunk_texts(
        self,
        texts: List[str],
        metadatas: Optional[List[DocumentMetadata]] = None,
    ) -> ChunkedDocument:
        """Split multiple texts into chunks."""
        all_chunks: List[str] = []
        all_meta: List[DocumentMetadata] = []
        for i, text in enumerate(texts):
            meta = metadatas[i] if metadatas and i < len(metadatas) else DocumentMetadata()
            chunks = self.splitter.split_text(text)
            all_chunks.extend(chunks)
            all_meta.extend([meta] * len(chunks))
        return ChunkedDocument(
            chunks=all_chunks, metadata=DocumentMetadata(), chunk_metadata=all_meta,
        )

    def _load_pdf(self, path: Path, base_meta: DocumentMetadata) -> ChunkedDocument:
        """Load PDF using pypdf (lazy import)."""
        try:
            from pypdf import PdfReader  # lazy import
        except ImportError:
            raise ImportError("pypdf required. Install: pip install pypdf") from None

        reader = PdfReader(str(path))
        all_text: List[str] = []
        all_meta: List[DocumentMetadata] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text.strip():
                page_meta = DocumentMetadata(
                    source=base_meta.source,
                    file_type="pdf",
                    page=page_num,
                )
                chunks = self.splitter.split_text(text)
                all_text.extend(chunks)
                all_meta.extend([page_meta] * len(chunks))

        logger.debug(
            "Loaded PDF %s: %d chunks from %d pages",
            path.name, len(all_text), len(reader.pages),
        )
        return ChunkedDocument(
            chunks=all_text,
            metadata=base_meta,
            chunk_metadata=all_meta,
        )

    def _load_docx(self, path: Path, base_meta: DocumentMetadata) -> ChunkedDocument:
        """Load DOCX using python-docx (lazy import)."""
        try:
            from docx import Document  # lazy import
        except ImportError:
            raise ImportError("python-docx required. Install: pip install python-docx") from None

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        chunks = self.splitter.split_text(text)
        meta_list = [base_meta] * len(chunks)

        logger.debug(
            "Loaded DOCX %s: %d chunks from %d paragraphs",
            path.name, len(chunks), len(paragraphs),
        )
        return ChunkedDocument(chunks=chunks, metadata=base_meta, chunk_metadata=meta_list)

    def _load_text(self, path: Path, base_meta: DocumentMetadata) -> ChunkedDocument:
        """Load TXT or MD file."""
        text = path.read_text(encoding="utf-8")
        chunks = self.splitter.split_text(text)
        meta_list = [base_meta] * len(chunks)

        logger.debug("Loaded text file %s: %d chunks", path.name, len(chunks))
        return ChunkedDocument(chunks=chunks, metadata=base_meta, chunk_metadata=meta_list)

    def load_text_from_upload(self, content: bytes, filename: str) -> ChunkedDocument:
        """
        Load from uploaded file bytes (saves to temp, processes, cleans up).
        """
        ext = Path(filename).suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return self.load_and_chunk(tmp_path, metadata={"source": filename})
        finally:
            os.unlink(tmp_path)
