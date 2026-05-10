"""
Shared fixtures and test configuration.
"""

from __future__ import annotations

import os
import tempfile
from typing import Generator

import pytest
from fastapi import FastAPI

from config import Config
from models.schemas import DocumentMetadata, SourceDocument
from services.cache_service import ResponseCache
from services.document_service import DocumentProcessor, TextSplitter
from services.vector_store import VectorStoreService


# ─── Test Config ───────────────────────────────────────────────────


@pytest.fixture
def test_config() -> Config:
    """Minimal config for tests — mock mode with temp dirs."""
    return Config(
        llm_provider="mock",
        mock_mode=True,
        chroma_db_dir=tempfile.mkdtemp(prefix="chroma_test_"),
        upload_dir=tempfile.mkdtemp(prefix="upload_test_"),
        debug=False,
    )


# ─── Temp Directory ────────────────────────────────────────────────


@pytest.fixture
def temp_dir() -> Generator[str, None, None]:
    """Create and clean up a temporary directory."""
    d = tempfile.mkdtemp(prefix="rag_test_")
    yield d
    import shutil
    shutil.rmtree(d, ignore_errors=True)


# ─── Vector Store ─────────────────────────────────────────────────


@pytest.fixture
def vector_store(temp_dir) -> VectorStoreService:
    """Create a test ChromaDB vector store with mock embeddings."""
    from services.embedding_service import MockEmbeddingService

    embed = MockEmbeddingService()
    store = VectorStoreService(
        persist_directory=os.path.join(temp_dir, "chroma"),
        embedding_service=embed,
        collection_name="test_collection",
    )
    return store


@pytest.fixture
def vector_store_no_embed(temp_dir) -> VectorStoreService:
    """Vector store without embedding service (text-based search)."""
    store = VectorStoreService(
        persist_directory=os.path.join(temp_dir, "chroma_no_embed"),
        embedding_service=None,
        collection_name="test_no_embed",
    )
    return store


# ─── Document Processor ───────────────────────────────────────────


@pytest.fixture
def doc_processor() -> DocumentProcessor:
    """Document processor with small chunks for testing."""
    return DocumentProcessor(chunk_size=100, chunk_overlap=10)


@pytest.fixture
def text_splitter() -> TextSplitter:
    """Text splitter with small chunks."""
    return TextSplitter(chunk_size=50, chunk_overlap=10)


# ─── Cache ─────────────────────────────────────────────────────────


@pytest.fixture
def cache() -> ResponseCache:
    """In-memory cache with short TTL for testing."""
    return ResponseCache(ttl_seconds=3600, max_size=16)


# ─── Sample Documents ──────────────────────────────────────────────


@pytest.fixture
def sample_text() -> str:
    return (
        "Наша компания是一家 ведущая IT-компания, основанная в 2010 году. "
        "Мы специализируемся на разработке программного обеспечения, "
        "искусственного интеллекта и облачных решений. "
        "Наш офис находится в Москве, также у нас есть представительства "
        "в Санкт-Петербурге и Казани. "
        "Мы обслуживаем более 500 корпоративных клиентов по всей России."
    )


@pytest.fixture
def sample_pdf_path(temp_dir) -> str:
    """Create a minimal valid PDF for testing."""
    pdf_path = os.path.join(temp_dir, "test_doc.pdf")
    # Minimal PDF content
    pdf_content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Resources<<>>>>endobj\n"
        b"xref\n"
        b"0 4\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"trailer<</Size 4/Root 1 0 R>>\n"
        b"startxref\n"
        b"190\n"
        b"%%EOF"
    )
    with open(pdf_path, "wb") as f:
        f.write(pdf_content)
    return pdf_path


@pytest.fixture
def sample_txt_path(temp_dir) -> str:
    """Create a sample text file."""
    txt_path = os.path.join(temp_dir, "test_doc.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(
            "Это тестовый документ для проверки RAG-системы.\n\n"
            "Он содержит несколько абзацев текста на русском языке.\n\n"
            "Компания ООО «Ромашка» была основана в 2015 году.\n"
            "Основные направления деятельности: консалтинг и разработка ПО.\n\n"
            "Контакты: г. Москва, ул. Ленина, д. 10, офис 501.\n"
            "Телефон: +7 (495) 123-45-67.\n"
        )
    return txt_path


@pytest.fixture
def sample_docx_path(temp_dir) -> str:
    """Create a minimal DOCX for testing (requires python-docx)."""
    try:
        from docx import Document
        docx_path = os.path.join(temp_dir, "test_doc.docx")
        doc = Document()
        doc.add_paragraph("Тестовый DOCX документ для RAG.")
        doc.add_paragraph("Он содержит информацию о политике конфиденциальности.")
        doc.save(docx_path)
        return docx_path
    except ImportError:
        pytest.skip("python-docx not installed")


# ─── Source Documents ──────────────────────────────────────────────


@pytest.fixture
def sample_source_docs() -> list[SourceDocument]:
    """Sample retrieved documents for testing LLM calls."""
    return [
        SourceDocument(
            content="Наша компания предоставляет услуги по разработке ИИ.",
            metadata=DocumentMetadata(source="about_us.txt", file_type="txt"),
            score=0.95,
        ),
        SourceDocument(
            content="Цены на услуги начинаются от 50 000 рублей.",
            metadata=DocumentMetadata(source="pricing.txt", file_type="txt"),
            score=0.82,
        ),
    ]
